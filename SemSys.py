#!/usr/bin/env python
# coding: utf-8

# In[1]:


import itertools
import sys
import os
from docx.api import Document
import re
import stanza
from docx import Document
import pathlib


# In[2]:


from docx import Document
import pathlib


# In[3]:


pathMetodichka = pathlib.PurePath(
    '/Users/simple/Metod_09.03.01_MOVTAS_SII_LR.docx')


# In[4]:


# присваиваем переменной название документа
metodichkaName = pathMetodichka.name

# преобразовываем в строковый формат
pathMetodichka = str(pathMetodichka)

# путь к документу для парсинга
document = Document(pathMetodichka)


# In[5]:


nlp = stanza.Pipeline('ru')


# In[6]:


# Парсинг абзацев документа по ключевым словам
key_word_theme = 'Тема работы'
key_word_goal = 'Цель работы'
key_word_description = 'Описание работы'


# In[7]:


def getSentences():
   # print(str(key_word_theme))
    for paragraph in document.paragraphs:
        if key_word_theme in paragraph.text:
            result_theme = paragraph.text.partition(key_word_theme)[2][2:]
            yield(result_theme)

    for paragraph in document.paragraphs:
        if key_word_goal in paragraph.text:
            result_goal = paragraph.text.partition(key_word_goal)[2][2:]
            yield(result_goal)

    for paragraph in document.paragraphs:
        if key_word_description in paragraph.text:
            result_description = paragraph.text.partition(
                key_word_description)[2][2:]
            yield(result_description)


# In[8]:

gen = getSentences()
doc = ("\n".join(map(str, gen)))

# Поиск Класса Дисциплина и его экземпляров
disciplina_key_word = "дисциплина"


# In[9]:


def getDiscip():
    # Устанавливаем лимит Параграфов для поиска ключевого слова в
    for paragraph in document.paragraphs[:]:
        if disciplina_key_word in paragraph.text:
            discip = paragraph.text.partition(disciplina_key_word)[2][2:]
            yield(discip)


foo = getDiscip()
disciplines = ("\n".join(map(str, foo)).replace(
    "«", " ").replace("»", " ").replace(" ", "_").split("."))
# первое предложение из абзаца
disciplines = disciplines[0]
disciplines = re.sub(r'\(.*\)', '', disciplines).split(",")


# In[10]:


# Выбираем таблицу с необходимой информацией[2] главная инфо про методичку
table = document.tables[2]


# In[11]:


# Создаем два генератора для парсинга левого и правого столбца
def firstColumn() -> dict:
    for stroka in (table.rows):
        column = stroka.cells[0].text
        yield(column)


def secondColumn():
    for stroka in (table.rows):
        column = (stroka.cells[1].text)

        # Удаляем  все сокращения в скобках регулярным выражением
        column = re.sub(r'\(.*\)', '', column)
        yield(column)


# Результат генератора передаем в список и присваиваем его переменной
g = list(firstColumn())
# Конвертируем список в строку для дальнейшей работы с форматом String
titul = ("\n".join(map(str, g)))
# print(titul)
g1 = list(secondColumn())
titulValues = ("".join(map(str, g1)).replace(" ", "_").split("."))
# for line in titulValues:
#   print(":{0} rdf:type owl:NamedIndividual, {0} .".format(line))
# f.write(": {0} rdf:type owl:NamedIndividual, {0} .".format(line))


# In[12]:


# Выбираем таблицу с необходимой информацией; [3] ФИО авторов
table = document.tables[3]

# Создаем два генератора для парсинга левого и правого столбца


def firstColumnFIO():
    for stroka in (table.rows):
        column = stroka.cells[0].text
        yield(column)


def secondColumnFIO():
    for stroka in (table.rows):
        column = (stroka.cells[1].text)

        # Удаляем  все сокращения в скобках регулярным выражением
        column = re.sub(r'\(.*\)', '', column)
        yield(column)


# Результат генератора передаем в список и присваиваем его переменной
genFCF = list(firstColumnFIO())

# Конвертируем список в строку для дальнейшей работы с форматом String
titulFIO = ("\n".join(map(str, genFCF)))
# print(titulFIO)

genSCF = list(secondColumnFIO())
titulValuesFIO = ("".join(map(str, genSCF)).replace(" ", "_", )).split(",")
# f.write(": {0} rdf:type owl:NamedIndividual, {0} .".format(line))


# In[13]:


doc = nlp(doc)
titul = nlp(titul)
titulFIO = nlp(titulFIO)


# In[14]:


# Вывод работы всех процессов анализа  парагаграфов
for i, sent in enumerate(doc.sentences):
    print("[Предложение {}]".format(i+1))
    for word in sent.words:
        print("{:20s}\t{:20s}\t{:6s}\t{:d}\t{:12s}".format(
            word.text, word.lemma, word.pos, word.head, word.deprel))
    print("")


# АНАЛИЗ ПРАВОЙ ЧАСТИ В ТАБЛИЦЕ2 НЕ ТРЕБУЕТСЯ БЕРЕМ ВСЕ КАК ЭКЗЕМПЛЯРЫ


# переменные для основных Classes
avtor = "Автор"
deystvia = "Действия"
obj = "Объект"
objPredObl = "Объект_предметной_области"


def unique_list(l):
    ulist = []
    [ulist.append(x) for x in l if x not in ulist]
    return ulist


# каждый раз перезаписваем файл
with open(sys.argv[1] + ".ttl", "w") as f:
    f.write("""
    @prefix : <http://www.semanticweb.org/simple/ontologies/2020/5/untitled-ontology-25#> .
    @prefix owl: <http://www.w3.org/2002/07/owl#> .
    @prefix rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#> .
    @prefix xml: <http://www.w3.org/XML/1998/namespace> .
    @prefix xsd: <http://www.w3.org/2001/XMLSchema#> .
    @prefix rdfs: <http://www.w3.org/2000/01/rdf-schema#> .
    <http://www.semanticweb.org/simple/ontologies/2020/5/system> rdf:type owl:Ontology . \n
    """)
    # ---------------------------------CLASSES---------------------------------

    # Создаем необходимые Classes статически, так как они необходимы для удобного разграничения ролей сущностей.
    f.write(""":{} rdf:type owl:Class .\n""".format(obj))  # Объект
    f.write(""":{} rdf:type owl:Class .\n""".format(
        objPredObl))  # Объект Предметной области

    # Запись в файл  левой части основного титульника как Classes и правой части как individuals
    titul_list = []
    for i, sent in enumerate(titul.sentences):
        for word in sent.words:
            if word.deprel == "root":
                titul_list.append(word.text)

    # Деструктуризация левого столбца с основной страницы
    naprav, profil, instit, formaObuch, progPodgot, kafedra = [
        str(e) for e in titul_list]

    print("")

    f.write(' '.join(f""":{tl} rdf:type owl:Class;
                        rdfs:subClassOf :Объект . \n""" for tl in (titul_list)))

    f.write(' '.join(f""":{tl} rdf:type owl:NamedIndividual,  
                              :{tv} . \n\n""" for tv, tl in zip(titul_list, titulValues)))

    # Анализ сущностей в  параграфе в документ
    import re
    import itertools
    stop_list = ["conj", "case"]
    para_list = []

    for i, sent in enumerate(doc.sentences):
        for word in sent.words:
            if word.deprel == "obl" and "PRON" not in word.pos:
                para_list.append(word.lemma)

            if word.deprel == "obj":
                para_list.append(word.lemma)

            if word.deprel == "iobj":
                para_list.append(word.lemma)

            if word.deprel == "nmod" and "PRON" not in word.pos:
                para_list.append(word.lemma)

            if word.deprel == "amod" and "VERB" not in word.pos:
                para_list.append(word.lemma)

            if word.deprel is not "amod":
                para_list.append(".")

    #s = ' '.join(unique_list(s.split()))

    s = ("".join(map(str, para_list)))
    #dot_pattern = re.compile(r'\.{1,}')
    #single_dot = dot_pattern.sub(' ', paraToStr)
    #''.join(i for i, _ in itertools.groupby(paraToStr))
    # print(single_dot)

    while '...' in s:
        s = s.replace('.......', '......')
        s = s.replace('......', '.....')
        s = s.replace('.....', '....')
        s = s.replace('....', '...')
        s = s.replace('...', '..')
        s = s.replace('..', '<temp>')
        s = s.replace('.', ' ')
        s = s.replace('<temp>', '.')
        s = s.replace(" ", "_")

    for i in s:
        s = re.sub('(^_*)', "", s)

    lines = s.split('.')
    for line in lines:

        f.write(""":{} rdf:type owl:Class;
                    rdfs:subClassOf :%s .\n""".format(line) % objPredObl)

    # root слова Существительные из текста как Classes
    s_roots = []
    for i, sent in enumerate(doc.sentences):
        for word in sent.words:
            if word.deprel == "root" and word.pos == "NOUN":
                s_roots.append(word.text)

    # Удалаяем дубликаты root слов в списке s_roots
    sNew = []
    for word in s_roots:
        if word not in sNew:
            sNew.append(word)
    s_roots = sNew

    listToStr = ' '.join([str(elem) for elem in s_roots]).split(" ")
    for word in listToStr:
        f.write(""":{} rdf:type owl:Class;
                    rdfs:subClassOf :%s . \n""".format(word) % deystvia)

    # Class "Методические указания" и "Автор"
    for i, sent in enumerate(titulFIO.sentences):
        for word in sent.words:
            if word.deprel == "amod" and word.head < 5:
                amod = word.text

            if word.deprel == "nsubj:pass":
                nsubj = word.text
                metod_ukaz = amod + " " + nsubj
                metod_ukaz = metod_ukaz.replace(" ", "_")
                f.write(""":{} rdf:type owl:Class;
                       rdfs:subClassOf :%s. \n\n""".format(metod_ukaz) % obj)
                f.write(""":{} rdf:type owl:Class;
                        rdfs:subClassOf :%s. \n\n""".format(avtor) % obj)
                break

    # Class "Дисциплина"
    for disc in disciplines:
        f.write(""":{} rdf:type owl:Class;
                        rdfs:subClassOf :%s. \n""".format(disciplina_key_word) % obj)

    # ---------------------------------OBJECT PROPERTIES---------------------------------

        # "методические указания" РАЗРАБОТАНЫ "Автор"
    for word in sent.words:
        if word.pos == "VERB":
            razrabotany = word.text
            f.write(""":{0} rdf:type owl:ObjectProperty ;
                                 rdfs:domain :{1};
                                 rdfs:range :%s . \n\n""".format(razrabotany, metod_ukaz) % avtor)
            break
    print("")

    # "мет.указания" ОПИСЫВАЮТ "действия"
    f.write(""":описывают rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(metod_ukaz, deystvia))

    # "мет.указания" ОПИСЫВАЮТ "объектв предметной области"
    f.write(""":описывают rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(metod_ukaz, objPredObl))

    # "методические указания" ДЛЯ "дисциплина"

    f.write(""":для rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(metod_ukaz, disciplina_key_word))

    # "методические указания" ДЛЯ "профиля"
    f.write(""":для rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(metod_ukaz, profil))

    # "методические указания" ПО "дисциплина"
    f.write(""":по  rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(metod_ukaz, disciplina_key_word))

    # "методические указания" ПО "программа подготовки"
    f.write(""":по  rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(metod_ukaz, progPodgot))

    # "кафедра" ОСУЩЕСТВЛЯЕТ ПОДГОТОВКУ ПО "профиль"
    f.write(""":осуществляет_подготовку_по  rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(kafedra, profil))

    # "кафедра" ВХОДИТ В СОСТАВ "институт"
    f.write(""":входит_в_состав  rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(kafedra, instit))

    # "профиль" ВХОДИТ В СОСТАВ "направление"
    f.write(""":входит_в_состав rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(profil, naprav))
    # "дисциплина" ВХОДИТ В СОСТАВ "профиль"
    f.write(""":входит_в_состав rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(disciplina_key_word, profil))
    # инверсное отношение. "институт" ИМЕЕТ В СОСТАВЕ "кафедра"
    f.write(""":имеет_в_составе rdf:type owl:ObjectProperty ;
                        rdfs:domain :{0};
                        rdfs:range :{1}. \n""".format(instit, kafedra))

    # Глаголы и Прилагательные из текста как ObjectProperty
    verb_adj = []
    for i, sent in enumerate(doc.sentences):
        for word in sent.words:
            if word.deprel == "root" and (word.pos == "VERB" or word.pos == "ADJ"):
                verb_adj.append(word.text)
            if word.deprel == "conj" and (word.pos == "VERB"):
                verb_adj.append(word.text)

    clear_dubl_verb_adj = []
    for word in verb_adj:
        if word not in clear_dubl_verb_adj:
            clear_dubl_verb_adj.append(word)
    verb_adj = clear_dubl_verb_adj

    for objProp in verb_adj:
        f.write(""":{} rdf:type owl:ObjectProperty .\n""".format(objProp))

    print("")
    # ---------------------------------INDIVIDUALS---------------------------------

    # Экземпляры класса Дисциплина
    for disc in disciplines:
        f.write(""":{} rdf:type owl:NamedIndividual,
                            :%s . \n""".format(disc) % disciplina_key_word)

    # Экземпляры класса Автор (ФИО)
    for line in titulValuesFIO:
        f.write(""":{0} rdf:type owl:NamedIndividual,
                            :%s . \n\n""".format(line) % avtor)

    # Экземпляр класса Методические указания
    for line in titulValuesFIO:
        f.write(""":{0} rdf:type owl:NamedIndividual,
                            :{1} ;
                    :{2} :%s .\n\n""".format(metodichkaName, metod_ukaz, razrabotany) % line)


print("Создание онтологии завершено")


# In[ ]:
